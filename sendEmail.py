import os
from email.mime.text import MIMEText
import docx

def update_dictionary(data):
	merged_data = {}

	for item in data:
		for key, value in item.items():
			merged_data.setdefault(key, []).append(value)

	result = {}

	for key, value in merged_data.items():
		result[key] = '\n'.join(value)

	return result



def create_docx(shablon, name, Dictionary, data, data1):
	Dictionary.update(update_dictionary(data))
	Dictionary.update(update_dictionary(data1))
	for key in Dictionary:
		Dictionary[key] = str(Dictionary[key])

	print(shablon)

	doc = docx.Document(shablon)
	style = doc.styles['Normal']
	font = style.font
	font.name = 'Times New Roman'
	font.size = docx.shared.Pt(12)

	for i in Dictionary:
		for p in doc.paragraphs:
			if p.text.find(i) >= 0:
				p.text = p.text.replace(i, str(Dictionary[i]))


	for j in Dictionary:
		for table in doc.tables:
			for col in table.columns:
				for cell in col.cells:
					for paragraph in cell.paragraphs:
						if j in paragraph.text:
							for run in paragraph.runs:
								if j in run.text:
									run.font.italic = True
									run.font.bold = True
									run.text = run.text.replace(j, str(Dictionary[j]))

	doc.save(name+'.docx')



import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

from transliterate import translit

def transliterate_string(text):
	# Выполняем транслитерацию с помощью библиотеки transliterate
	transliterated_text = translit(text, 'ru', reversed=True)
	return transliterated_text

def send_email(sender_email, sender_password, receiver_email, subject, body, file_path):
	file_name = os.path.basename(file_path)
	file_name = transliterate_string(file_name)
	# Создание объекта MIMEMultipart
	message = MIMEMultipart()
	message["From"] = sender_email
	message["To"] = receiver_email
	message["Subject"] = subject

	# Добавление текстового содержимого письма
	message.attach(MIMEText(body, "plain"))

	# Открытие и добавление файла в письмо
	with open(file_path, "rb") as attachment:
		part = MIMEBase("application", "octet-stream")
		part.set_payload(attachment.read())

	# Кодирование файла в формат Base64
	encoders.encode_base64(part)

	# Добавление заголовков файла
	part.add_header(
		"Content-Disposition",
		f"attachment; filename= {file_name}",
	)

	# Прикрепление файла к письму
	message.attach(part)

	# Отправка письма через SMTP-сервер
	with smtplib.SMTP("smtp.gmail.com", 587) as server:
		server.starttls()
		server.login(sender_email, sender_password)
		server.send_message(message)

	os.remove(file_path)


